"""
DeepFind Engine — DOCX Text Extractor

Uses python-docx for local, offline Word document text extraction.
No cloud APIs. No OCR.

Handles:
  - Normal .docx files
  - Corrupted/invalid .docx files (skips gracefully)

Note: .doc (old binary format) is not supported — only .docx.
"""

import logging
from pathlib import Path

from extractors.base_extractor import BaseExtractor, ExtractorError

log = logging.getLogger(__name__)


class DocxExtractor(BaseExtractor):
    """
    Extracts text from .docx files using python-docx.

    Extracts text from all paragraphs in order.
    Tables are partially supported (text from each cell).
    Headers and footers are not extracted (out of Step 9 scope).
    """

    def _extract_raw(self, path: Path) -> str:
        try:
            import docx  # type: ignore  (python-docx)
        except ImportError:
            raise ExtractorError(
                "python-docx is not installed. Run: pip install python-docx"
            )

        try:
            doc = docx.Document(str(path))
        except Exception as exc:
            raise ExtractorError(f"DOCX open error: {exc}") from exc

        parts: list[str] = []

        # Main body paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Table cell text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        parts.append(text)

        return "\n".join(parts)
